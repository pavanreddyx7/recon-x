#!/usr/bin/env python3
"""Convert PROJECT_REPORT.md to PROJECT_REPORT.docx with proper academic formatting."""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH   = os.path.join(os.path.dirname(__file__), "PROJECT_REPORT.md")
DOCX_PATH = os.path.join(os.path.dirname(__file__), "PROJECT_REPORT.docx")


# ── helpers ──────────────────────────────────────────────────────────────────

def set_paragraph_spacing(para, before=0, after=6, line_rule=WD_LINE_SPACING.ONE_POINT_FIVE):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing_rule = line_rule


def add_page_break(doc):
    para = doc.add_paragraph()
    run  = para.add_run()
    run.add_break(__import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)


def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def apply_inline(run_parent, text):
    """Add runs with bold/italic/code inline formatting parsed from text."""
    # patterns: **bold**, *italic*, `code`
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    pos = 0
    for m in pattern.finditer(text):
        # plain text before match
        if m.start() > pos:
            run_parent.add_run(text[pos:m.start()])
        full = m.group(0)
        if full.startswith("**"):
            r = run_parent.add_run(m.group(2))
            r.bold = True
        elif full.startswith("`"):
            r = run_parent.add_run(m.group(4))
            r.font.name = "Courier New"
            r.font.size = Pt(10)
        else:
            r = run_parent.add_run(m.group(3))
            r.italic = True
        pos = m.end()
    if pos < len(text):
        run_parent.add_run(text[pos:])


def add_styled_paragraph(doc, text, style="Normal", bold=False, italic=False,
                          align=WD_ALIGN_PARAGRAPH.LEFT, size=None, color=None,
                          before=0, after=6):
    para = doc.add_paragraph(style=style)
    apply_inline(para, text)
    for run in para.runs:
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    para.alignment = align
    set_paragraph_spacing(para, before=before, after=after)
    return para


def add_blank(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, 0, 0)


def render_md_table(doc, header_row, rows):
    col_count = len(header_row)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header
    hrow = table.rows[0]
    for i, cell_text in enumerate(header_row):
        cell = hrow.cells[i]
        set_cell_bg(cell, "1F3864")
        para = cell.paragraphs[0]
        run  = para.add_run(cell_text.strip().strip("*"))
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # data rows
    for r_idx, row_data in enumerate(rows):
        row_obj = table.rows[r_idx + 1]
        bg = "EBF3FB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            cell = row_obj.cells[c_idx]
            set_cell_bg(cell, bg)
            para = cell.paragraphs[0]
            text = cell_text.strip()
            # center-align columns that are numeric/symbol heavy
            if re.match(r"^[\d✓✗⚠\.\-\+\s/]+$", text):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_inline(para, text)
            for run in para.runs:
                run.font.size = Pt(10)

    doc.add_paragraph()


# ── main converter ────────────────────────────────────────────────────────────

def convert():
    with open(MD_PATH, encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)

    # Default font
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    # Heading styles
    for lvl, sz in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)]:
        s = doc.styles[lvl]
        s.font.name = "Times New Roman"
        s.font.size = Pt(sz)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    i = 0
    in_table    = False
    table_lines = []
    in_code     = False
    code_lines  = []

    page_break_next = False

    while i < len(lines):
        raw  = lines[i]
        line = raw.rstrip("\n")
        stripped = line.strip()

        # ── code block ────────────────────────────────────────────────────
        if stripped.startswith("```"):
            if not in_code:
                in_code    = True
                code_lines = []
            else:
                in_code = False
                code_para = doc.add_paragraph()
                code_para.style = "Normal"
                run = code_para.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                set_paragraph_spacing(code_para, 4, 4)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── markdown table ─────────────────────────────────────────────────
        if stripped.startswith("|"):
            if not in_table:
                in_table    = True
                table_lines = []
            table_lines.append(stripped)
            i += 1
            continue
        else:
            if in_table:
                in_table = False
                # parse table_lines
                data_rows = [l for l in table_lines
                             if not re.match(r"^\|[-:\s|]+\|?$", l)]
                if data_rows:
                    header = [c for c in data_rows[0].split("|") if c.strip() != ""]
                    rows   = [[c for c in r.split("|") if c.strip() != ""]
                               for r in data_rows[1:]]
                    render_md_table(doc, header, rows)
                table_lines = []

        # ── horizontal rule ────────────────────────────────────────────────
        if re.match(r"^---+$", stripped):
            if page_break_next:
                add_page_break(doc)
                page_break_next = False
            else:
                page_break_next = True
            i += 1
            continue

        # ── blank / nbsp ───────────────────────────────────────────────────
        if stripped == "" or stripped == "&nbsp;":
            add_blank(doc)
            i += 1
            continue

        # ── headings ───────────────────────────────────────────────────────
        h1 = re.match(r"^# (.+)$", line)
        h2 = re.match(r"^## (.+)$", line)
        h3 = re.match(r"^### (.+)$", line)
        h4 = re.match(r"^#### (.+)$", line)

        if h1:
            text = h1.group(1).strip()
            if page_break_next:
                add_page_break(doc)
                page_break_next = False
            para = doc.add_heading(text, level=1)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(para, 12, 6)
            i += 1
            continue

        if h2:
            text = h2.group(1).strip()
            if page_break_next:
                add_page_break(doc)
                page_break_next = False
            para = doc.add_heading(text, level=2)
            set_paragraph_spacing(para, 18, 6)
            i += 1
            continue

        if h3:
            para = doc.add_heading(h3.group(1).strip(), level=3)
            set_paragraph_spacing(para, 12, 4)
            i += 1
            continue

        if h4:
            para = doc.add_heading(h4.group(1).strip(), level=4)
            set_paragraph_spacing(para, 8, 4)
            i += 1
            continue

        # ── ordered list ───────────────────────────────────────────────────
        ol = re.match(r"^(\d+)\. (.+)$", stripped)
        if ol:
            para = doc.add_paragraph(style="List Number")
            apply_inline(para, ol.group(2))
            for run in para.runs:
                run.font.size = Pt(12)
            set_paragraph_spacing(para, 0, 3)
            i += 1
            continue

        # ── unordered list ─────────────────────────────────────────────────
        ul = re.match(r"^[-*•] (.+)$", stripped)
        if ul:
            para = doc.add_paragraph(style="List Bullet")
            apply_inline(para, ul.group(1))
            for run in para.runs:
                run.font.size = Pt(12)
            set_paragraph_spacing(para, 0, 3)
            i += 1
            continue

        # ── italic/placeholder lines (*[Figure...]*) ───────────────────────
        if stripped.startswith("*[") and stripped.endswith("]*"):
            inner = stripped[1:-1]   # strip outer *
            para  = doc.add_paragraph()
            run   = para.add_run(inner)
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(para, 6, 6)
            i += 1
            continue

        # ── placeholder/note lines (lines starting with *( and ending )*) ──
        if stripped.startswith("*(") and stripped.endswith(")*"):
            para = doc.add_paragraph()
            run  = para.add_run(stripped[2:-2])
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(para, 2, 2)
            i += 1
            continue

        # ── normal paragraph ───────────────────────────────────────────────
        if stripped:
            if page_break_next:
                add_page_break(doc)
                page_break_next = False
            para = doc.add_paragraph()
            apply_inline(para, stripped)
            for run in para.runs:
                run.font.size = Pt(12)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(para, 0, 6)

        i += 1

    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    convert()
